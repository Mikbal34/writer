"""
BookDocument - Generalized DOCX builder.
Ported from TezDocument with configurable fonts, margins, and line spacing.
Supports real Word footnotes via XML manipulation, quotes, headings, and bibliography.
"""

import os
import re
import zipfile

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
FOOTNOTES_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
)

DEFAULT_FONTS = {
    "main": "Times New Roman",
    "size_title": 16,
    "size_h1": 14,
    "size_h2": 12,
    "size_h3": 12,
    "size_h4": 12,
    "size_main": 12,
    "size_block_quote": 11,
    "size_footnote": 10,
    "size_bibliography": 11,
}

DEFAULT_MARGINS = {
    "top": 2.5,
    "bottom": 2.5,
    "left": 3.5,
    "right": 2.5,
}


class BookDocument:
    """
    Generalized document builder with real Word footnotes,
    configurable styling, headings with auto-numbering,
    quote systems (inline, block, translated), and bibliography support.
    """

    def __init__(
        self,
        fonts: dict | None = None,
        margins: dict | None = None,
        line_spacing: float = 1.5,
        output_dir: str = ".",
    ):
        self.doc = Document()
        self.footnote_id = 1
        self.footnotes_data: list[tuple[int, str]] = []
        self.heading_counters = [0, 0, 0, 0]

        self.fonts = {**DEFAULT_FONTS, **(fonts or {})}
        self.margins = {**DEFAULT_MARGINS, **(margins or {})}
        self.line_spacing = line_spacing
        self.output_dir = output_dir

        # Apply page margins
        for section in self.doc.sections:
            section.top_margin = Cm(self.margins["top"])
            section.bottom_margin = Cm(self.margins["bottom"])
            section.left_margin = Cm(self.margins["left"])
            section.right_margin = Cm(self.margins["right"])

    # ===================== Font Helpers =====================

    def _set_font(
        self,
        run,
        font_name: str | None = None,
        size: int | None = None,
        bold: bool = False,
        italic: bool = False,
    ):
        """Apply font settings to a run."""
        if font_name is None:
            font_name = self.fonts["main"]
        if size is None:
            size = self.fonts["size_main"]

        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # ===================== Heading System =====================

    def add_book_title(self, text: str):
        """Add a centered book/document title."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(text)
        self._set_font(run, size=self.fonts["size_title"], bold=True)
        return p

    def add_heading_level1(self, title: str, auto_number: bool = True):
        """Level 1 heading (e.g. '1. Chapter Title')."""
        if auto_number:
            self.heading_counters[0] += 1
            self.heading_counters[1] = 0
            self.heading_counters[2] = 0
            self.heading_counters[3] = 0
            full_title = f"{self.heading_counters[0]}. {title}"
        else:
            full_title = title

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(full_title)
        self._set_font(run, size=self.fonts["size_h1"], bold=True)
        return p

    def add_heading_level2(self, title: str, auto_number: bool = True):
        """Level 2 heading (e.g. '1.1. Section Title')."""
        if auto_number:
            self.heading_counters[1] += 1
            self.heading_counters[2] = 0
            self.heading_counters[3] = 0
            number = f"{self.heading_counters[0]}.{self.heading_counters[1]}"
            full_title = f"{number}. {title}"
        else:
            full_title = title

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(full_title)
        self._set_font(run, size=self.fonts["size_h2"], bold=True)
        return p

    def add_heading_level3(self, title: str, auto_number: bool = True):
        """Level 3 heading (e.g. '1.1.1. Subsection Title')."""
        if auto_number:
            self.heading_counters[2] += 1
            self.heading_counters[3] = 0
            number = (
                f"{self.heading_counters[0]}.{self.heading_counters[1]}"
                f".{self.heading_counters[2]}"
            )
            full_title = f"{number}. {title}"
        else:
            full_title = title

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(full_title)
        self._set_font(run, size=self.fonts["size_h3"], bold=True)
        return p

    def add_heading_level4(self, title: str, auto_number: bool = True):
        """Level 4 heading (e.g. '1.1.1.1. Sub-subsection Title')."""
        if auto_number:
            self.heading_counters[3] += 1
            number = (
                f"{self.heading_counters[0]}.{self.heading_counters[1]}"
                f".{self.heading_counters[2]}.{self.heading_counters[3]}"
            )
            full_title = f"{number}. {title}"
        else:
            full_title = title

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(full_title)
        self._set_font(run, size=self.fonts["size_h4"], bold=True)
        return p

    def reset_heading_counters(self):
        """Reset all heading counters to zero."""
        self.heading_counters = [0, 0, 0, 0]

    # ===================== Paragraphs =====================

    def add_paragraph(self, text: str, first_indent: bool = True):
        """
        Add a normal paragraph.
        Justified, with configurable line spacing and optional first-line indent.
        """
        p = self.doc.add_paragraph()
        if first_indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run(text)
        self._set_font(run)
        return p

    def add_paragraph_no_indent(self, text: str):
        """Add a paragraph without first-line indent."""
        return self.add_paragraph(text, first_indent=False)

    # ===================== Quote System =====================

    def add_inline_quote(self, text_before: str, quote: str, text_after: str = ""):
        """
        Short inline quote with double quotation marks.
        For quotes under 3 lines / 40 words.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if text_before:
            run1 = p.add_run(text_before)
            self._set_font(run1)

        run_quote = p.add_run(f'\u201c{quote}\u201d')
        self._set_font(run_quote)

        if text_after:
            run2 = p.add_run(text_after)
            self._set_font(run2)

        return p

    def add_nested_quote(
        self,
        text_before: str,
        outer_quote: str,
        inner_quote: str,
        text_after: str = "",
    ):
        """
        Nested quote. Outer: double quotes, Inner: single quotes.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if text_before:
            run1 = p.add_run(text_before)
            self._set_font(run1)

        full_quote = outer_quote.replace(inner_quote, f"\u2018{inner_quote}\u2019")
        run_quote = p.add_run(f"\u201c{full_quote}\u201d")
        self._set_font(run_quote)

        if text_after:
            run2 = p.add_run(text_after)
            self._set_font(run2)

        return p

    def add_block_quote(self, text: str):
        """
        Long block quote. Reduced font, single line spacing, left indent, no quotes.
        For quotes of 3+ lines / 40+ words.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run(text)
        self._set_font(run, size=self.fonts["size_block_quote"])
        return p

    def add_translated_quote(
        self, text_before: str, translation: str, text_after: str = ""
    ):
        """
        Translated quote from classical sources, rendered in italic.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if text_before:
            run1 = p.add_run(text_before)
            self._set_font(run1)

        run_translation = p.add_run(f'\u201c{translation}\u201d')
        self._set_font(run_translation, italic=True)

        if text_after:
            run2 = p.add_run(text_after)
            self._set_font(run2)

        return p

    @staticmethod
    def omission() -> str:
        """Omission marker for quotes: (...)"""
        return "(...)"

    # ===================== Italic Helpers =====================

    def add_text_with_italic(self, paragraph, text: str, italic_text: str):
        """Add normal text followed by italic text to an existing paragraph."""
        if text:
            run1 = paragraph.add_run(text)
            self._set_font(run1)

        run_italic = paragraph.add_run(italic_text)
        self._set_font(run_italic, italic=True)
        return paragraph

    def add_paragraph_with_book_title(
        self, text_before: str, book_title: str, text_after: str = ""
    ):
        """Add a paragraph with an italic book title."""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if text_before:
            run1 = p.add_run(text_before)
            self._set_font(run1)

        run_book = p.add_run(book_title)
        self._set_font(run_book, italic=True)

        if text_after:
            run2 = p.add_run(text_after)
            self._set_font(run2)

        return p

    # ===================== Footnote System =====================

    def add_footnote(self, paragraph, footnote_text: str) -> int:
        """
        Add a real Word footnote (appears at bottom of page).
        Must be called after adding the paragraph content.
        Returns the footnote ID.
        """
        fn_id = self.footnote_id
        self.footnotes_data.append((fn_id, footnote_text))
        self.footnote_id += 1

        # Add footnoteReference to the paragraph's XML
        run = paragraph.add_run()
        r = run._r

        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "FootnoteReference")
        rPr.append(rStyle)
        r.insert(0, rPr)

        footnoteRef = OxmlElement("w:footnoteReference")
        footnoteRef.set(qn("w:id"), str(fn_id))
        r.append(footnoteRef)

        return fn_id

    # ===================== Bibliography =====================

    def add_bibliography_title(self, title: str = "BIBLIOGRAPHY"):
        """Add a centered bibliography section title."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(title)
        self._set_font(run, size=self.fonts["size_h1"], bold=True)
        return p

    def add_bibliography_entry(self, text: str):
        """Add a bibliography entry with hanging indent."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(text)
        self._set_font(run, size=self.fonts["size_bibliography"])
        return p

    @staticmethod
    def sort_bibliography(entries: list[str]) -> list[str]:
        """Sort bibliography entries alphabetically, ignoring common prefixes."""
        def sort_key(entry: str) -> str:
            clean = re.sub(r"^(el-|er-|El-|Er-|al-|Al-|van |Von |de |De )", "", entry)
            return clean.lower()

        return sorted(entries, key=sort_key)

    # ===================== Cover / Front Matter =====================

    def add_cover_title(self, title: str, subtitle: str | None = None):
        """Add a centered cover page title with optional subtitle."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(120)
        run = p.add_run(title)
        self._set_font(run, size=18, bold=True)

        if subtitle:
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(subtitle)
            self._set_font(run2, size=14)

        return p

    def add_cover_author(self, author_name: str):
        """Add author name to cover page."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(48)
        run = p.add_run(author_name)
        self._set_font(run, size=14)
        return p

    def add_cover_info(self, info_text: str):
        """Add institutional info to cover page."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(72)
        run = p.add_run(info_text)
        self._set_font(run, size=12)
        return p

    def add_abstract_section(
        self, title: str, content: str, keywords: list[str]
    ):
        """Add an abstract/summary section with keywords."""
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(18)
        run_title = p_title.add_run(title)
        self._set_font(run_title, size=14, bold=True)

        p_content = self.doc.add_paragraph()
        p_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_content.paragraph_format.line_spacing = self.line_spacing
        run_content = p_content.add_run(content)
        self._set_font(run_content)

        p_keywords = self.doc.add_paragraph()
        p_keywords.paragraph_format.space_before = Pt(12)
        run_label = p_keywords.add_run("Keywords: ")
        self._set_font(run_label, bold=True)
        run_keys = p_keywords.add_run(", ".join(keywords))
        self._set_font(run_keys)

        return p_title, p_content, p_keywords

    # ===================== Table & Figure =====================

    def add_table_caption(self, number: int, title: str):
        """Add a table caption above a table."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"Table {number}: {title}")
        self._set_font(run, size=11, bold=True)
        return p

    def add_figure_caption(self, number: int, title: str):
        """Add a figure caption below a figure."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"Figure {number}: {title}")
        self._set_font(run, size=11, bold=True)
        return p

    def add_simple_table(self, headers: list[str], rows: list[list[str]]):
        """Add a simple table with headers and data rows."""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = self.fonts["main"]

        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                row.cells[i].text = str(cell_text)
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = self.fonts["main"]

        return table

    # ===================== Page / Section Breaks =====================

    def add_page_break(self):
        """Insert a page break."""
        self.doc.add_page_break()

    def add_section_break(self):
        """Insert a section break (for changing page number formats, etc.)."""
        p = self.doc.add_paragraph()
        sectPr = OxmlElement("w:sectPr")
        p._element.append(sectPr)
        return p

    # ===================== Save =====================

    def _create_footnotes_xml(self) -> bytes:
        """Build the footnotes.xml content."""
        nsmap_fn = {
            "w": WORD_NS,
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }

        footnotes = etree.Element("{%s}footnotes" % WORD_NS, nsmap=nsmap_fn)

        # Separator footnote (id=0)
        fn_sep = etree.SubElement(footnotes, "{%s}footnote" % WORD_NS)
        fn_sep.set("{%s}type" % WORD_NS, "separator")
        fn_sep.set("{%s}id" % WORD_NS, "0")
        p_sep = etree.SubElement(fn_sep, "{%s}p" % WORD_NS)
        r_sep = etree.SubElement(p_sep, "{%s}r" % WORD_NS)
        etree.SubElement(r_sep, "{%s}separator" % WORD_NS)

        # Continuation separator (id=-1)
        fn_cont = etree.SubElement(footnotes, "{%s}footnote" % WORD_NS)
        fn_cont.set("{%s}type" % WORD_NS, "continuationSeparator")
        fn_cont.set("{%s}id" % WORD_NS, "-1")
        p_cont = etree.SubElement(fn_cont, "{%s}p" % WORD_NS)
        r_cont = etree.SubElement(p_cont, "{%s}r" % WORD_NS)
        etree.SubElement(r_cont, "{%s}continuationSeparator" % WORD_NS)

        # Actual footnotes
        font_name = self.fonts["main"]
        fn_size = str(self.fonts["size_footnote"] * 2)  # half-points

        for fn_id, fn_text in self.footnotes_data:
            fn = etree.SubElement(footnotes, "{%s}footnote" % WORD_NS)
            fn.set("{%s}id" % WORD_NS, str(fn_id))

            p = etree.SubElement(fn, "{%s}p" % WORD_NS)

            # Paragraph properties
            pPr = etree.SubElement(p, "{%s}pPr" % WORD_NS)
            pStyle = etree.SubElement(pPr, "{%s}pStyle" % WORD_NS)
            pStyle.set("{%s}val" % WORD_NS, "FootnoteText")

            # Justify
            jc = etree.SubElement(pPr, "{%s}jc" % WORD_NS)
            jc.set("{%s}val" % WORD_NS, "both")

            # Single line spacing
            spacing = etree.SubElement(pPr, "{%s}spacing" % WORD_NS)
            spacing.set("{%s}line" % WORD_NS, "240")  # 240 twips = 1.0
            spacing.set("{%s}lineRule" % WORD_NS, "auto")
            spacing.set("{%s}before" % WORD_NS, "0")
            spacing.set("{%s}after" % WORD_NS, "0")

            # Footnote reference number
            r1 = etree.SubElement(p, "{%s}r" % WORD_NS)
            rPr1 = etree.SubElement(r1, "{%s}rPr" % WORD_NS)
            rStyle1 = etree.SubElement(rPr1, "{%s}rStyle" % WORD_NS)
            rStyle1.set("{%s}val" % WORD_NS, "FootnoteReference")
            etree.SubElement(r1, "{%s}footnoteRef" % WORD_NS)

            # Space after number
            r2 = etree.SubElement(p, "{%s}r" % WORD_NS)
            t2 = etree.SubElement(r2, "{%s}t" % WORD_NS)
            t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t2.text = " "

            # Footnote text
            r3 = etree.SubElement(p, "{%s}r" % WORD_NS)
            rPr3 = etree.SubElement(r3, "{%s}rPr" % WORD_NS)
            sz = etree.SubElement(rPr3, "{%s}sz" % WORD_NS)
            sz.set("{%s}val" % WORD_NS, fn_size)
            szCs = etree.SubElement(rPr3, "{%s}szCs" % WORD_NS)
            szCs.set("{%s}val" % WORD_NS, fn_size)
            rFonts = etree.SubElement(rPr3, "{%s}rFonts" % WORD_NS)
            rFonts.set("{%s}ascii" % WORD_NS, font_name)
            rFonts.set("{%s}hAnsi" % WORD_NS, font_name)
            t3 = etree.SubElement(r3, "{%s}t" % WORD_NS)
            t3.text = fn_text

        return etree.tostring(
            footnotes, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    def _update_rels(self, data: bytes) -> bytes:
        """Add footnotes relationship to document.xml.rels."""
        if not self.footnotes_data:
            return data

        root = etree.fromstring(data)
        existing_ids = [
            int(r.get("Id")[3:])
            for r in root
            if r.get("Id", "").startswith("rId")
        ]
        new_id = max(existing_ids) + 1 if existing_ids else 1

        rel = etree.SubElement(root, "Relationship")
        rel.set("Id", f"rId{new_id}")
        rel.set("Type", FOOTNOTES_REL_TYPE)
        rel.set("Target", "footnotes.xml")

        return etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    def _update_content_types(self, data: bytes) -> bytes:
        """Add footnotes override to [Content_Types].xml."""
        if not self.footnotes_data:
            return data

        root = etree.fromstring(data)
        ns = root.nsmap.get(
            None, "http://schemas.openxmlformats.org/package/2006/content-types"
        )

        override = etree.SubElement(root, "{%s}Override" % ns)
        override.set("PartName", "/word/footnotes.xml")
        override.set(
            "ContentType",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        )

        return etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    def save(self, filename: str, subfolder: str | None = None) -> str:
        """
        Save the document to disk, injecting footnotes XML.

        Args:
            filename: File name (with or without .docx extension).
            subfolder: Optional subfolder within output_dir.

        Returns:
            Full path to the saved file.
        """
        if subfolder:
            output_path = os.path.join(self.output_dir, subfolder)
        else:
            output_path = self.output_dir

        os.makedirs(output_path, exist_ok=True)

        if not filename.endswith(".docx"):
            filename += ".docx"
        filepath = os.path.join(output_path, filename)

        temp_path = filepath + ".temp"
        self.doc.save(temp_path)

        with zipfile.ZipFile(temp_path, "r") as zip_read:
            with zipfile.ZipFile(filepath, "w", zipfile.ZIP_DEFLATED) as zip_write:
                for item in zip_read.infolist():
                    data = zip_read.read(item.filename)

                    if item.filename == "word/_rels/document.xml.rels":
                        data = self._update_rels(data)
                    if item.filename == "[Content_Types].xml":
                        data = self._update_content_types(data)

                    zip_write.writestr(item, data)

                if self.footnotes_data:
                    footnotes_xml = self._create_footnotes_xml()
                    zip_write.writestr("word/footnotes.xml", footnotes_xml)

        os.remove(temp_path)
        return filepath


# ===================== Merge Functions =====================


def merge_documents(
    doc_paths: list[str],
    output_path: str,
    add_page_breaks: bool = True,
) -> str:
    """
    Merge multiple DOCX files into one.

    Args:
        doc_paths: Ordered list of file paths to merge.
        output_path: Path for the merged output file.
        add_page_breaks: Whether to insert page breaks between documents.

    Returns:
        Path to the merged file.
    """
    if not doc_paths:
        raise ValueError("At least one document path is required")

    merged = Document(doc_paths[0])

    for doc_path in doc_paths[1:]:
        if add_page_breaks:
            merged.add_page_break()

        sub_doc = Document(doc_path)
        for element in sub_doc.element.body:
            merged.element.body.append(element)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    merged.save(output_path)
    return output_path


def is_long_quote(text: str, line_threshold: int = 3, word_threshold: int = 40) -> bool:
    """Determine whether a quote should be block (True) or inline (False)."""
    lines = text.count("\n") + 1
    words = len(text.split())
    return lines >= line_threshold or words >= word_threshold

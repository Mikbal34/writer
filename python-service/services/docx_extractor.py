"""
DOCX text + metadata extraction.

A .docx has no real pages either — Word inserts page breaks at runtime
based on the user's printer/zoom. We extract paragraphs in order and
group them into ~25-paragraph synthetic "pages" so downstream chunking
+ citation UX has a natural position marker. The reader sees this as
a "konum" (position), and the citation displays as ¶ N.

Heading runs become their own short page so a lookup like
"section 3.1" tends to land on the right block.

Document `core_properties` (title / author / subject / created /
keywords) is also surfaced so the Node side can skip Haiku for
correctly-tagged Word files.
"""

from __future__ import annotations

import io

from docx import Document


# How many paragraphs we group into a single synthetic "page". Tuned
# loosely against academic prose where 25 paragraphs ≈ 1 print page.
_PARAGRAPHS_PER_PAGE = 25


def _looks_like_heading(p) -> bool:
    style = (p.style.name or "").lower() if p.style else ""
    return style.startswith("heading") or style == "title"


def extract_docx_pages(file_bytes: bytes) -> list[dict]:
    """Return [{ page_number, content }]. page_number is a 1-based
    synthetic position marker (paragraph block index)."""
    doc = Document(io.BytesIO(file_bytes))

    pages: list[dict] = []
    page_number = 0
    buffer: list[str] = []

    def flush() -> None:
        nonlocal page_number, buffer
        if not buffer:
            return
        page_number += 1
        pages.append({"page_number": page_number, "content": "\n\n".join(buffer)})
        buffer = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        if _looks_like_heading(p):
            # Headings start a fresh block so they're easy to locate.
            flush()
            buffer.append(text)
            flush()
            continue
        buffer.append(text)
        if len(buffer) >= _PARAGRAPHS_PER_PAGE:
            flush()
    flush()

    # Tables are a common place for academic content — fold their text
    # in too, treating each table as its own block.
    for tbl in doc.tables:
        cells: list[str] = []
        for row in tbl.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    cells.append(t)
        if cells:
            page_number += 1
            pages.append({"page_number": page_number, "content": "\n".join(cells)})

    return pages


def get_docx_total_pages(file_bytes: bytes) -> int:
    """Synthetic page count derived from paragraph blocks (see module docstring)."""
    return len(extract_docx_pages(file_bytes))


def extract_docx_metadata(file_bytes: bytes) -> dict:
    """Read the document's core_properties (set by Word's File →
    Properties dialog). All fields optional."""
    doc = Document(io.BytesIO(file_bytes))
    props = doc.core_properties
    out: dict = {}

    title = (props.title or "").strip() if props.title else ""
    if title:
        out["title"] = title

    author = (props.author or "").strip() if props.author else ""
    if author:
        out["author"] = author

    subject = (props.subject or "").strip() if props.subject else ""
    if subject:
        # The subject field is what people use for abstracts when the
        # document has no formal abstract section.
        out["abstract"] = subject

    keywords = (props.keywords or "").strip() if props.keywords else ""
    if keywords:
        # Word stores these as a comma- or semicolon-separated string.
        tokens = [t.strip() for t in keywords.replace(";", ",").split(",")]
        cleaned = [t for t in tokens if t]
        if cleaned:
            out["keywords"] = cleaned

    created = props.created
    if created and hasattr(created, "year"):
        out["year"] = str(created.year)

    return out
